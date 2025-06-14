import openpyxl
from openpyxl.styles import Font, Alignment
from openpyxl.drawing.image import Image
import requests
from io import BytesIO

def download_image(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        return BytesIO(response.content)
    except Exception as e:
        print(f"No se pudo descargar la imagen de {url}: {e}")
        return None

def create_excel_with_images():
    data = [
        ('Familia', 'Personaje Destacado', 'Siglo', 'Rol / Contribución', 'Imagen'),
        ('Romero', 'José Romero Silva', 'XVIII', 'Terrateniente, fundador de haciendas', 'https://upload.wikimedia.org/wikipedia/commons/thumb/2/2d/Coat_of_Arms_of_Spain.svg/120px-Coat_of_Arms_of_Spain.svg.png'),
        ('Hernández', 'Manuel Hernández Fuentes', 'XVIII', 'Agricultor destacado en Ñuble', 'https://upload.wikimedia.org/wikipedia/commons/thumb/a/a3/Coat_of_Arms_of_Chile.svg/120px-Coat_of_Arms_of_Chile.svg.png'),
        ('Valdés', 'José Miguel Valdés Hernández', 'XIX', 'Político, empresario agrícola', 'https://upload.wikimedia.org/wikipedia/commons/thumb/a/a3/Coat_of_Arms_of_Chile.svg/120px-Coat_of_Arms_of_Chile.svg.png'),
        # Añade más filas con URLs válidas para las imágenes
    ]

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Personajes Destacados"

    for row_index, row_data in enumerate(data, 1):
        for col_index, value in enumerate(row_data, 1):
            if row_index == 1 or col_index != 5:
                # Encabezados o datos normales
                cell = ws.cell(row=row_index, column=col_index, value=value)
                if row_index == 1:
                    cell.font = Font(bold=True)
                    cell.alignment = Alignment(horizontal='center')
                else:
                    cell.alignment = Alignment(wrap_text=True)
            else:
                # Insertar imagen en columna 5 (índice 4)
                img_data = download_image(value)
                if img_data:
                    img = Image(img_data)
                    # Ajustar tamaño imagen si quieres (ejemplo 50x50)
                    img.width = 50
                    img.height = 50
                    # Agregar imagen en la celda (columna 5, fila actual)
                    cell_location = f"E{row_index}"
                    ws.add_image(img, cell_location)
                else:
                    ws.cell(row=row_index, column=col_index, value="Imagen no disponible")

    # Ajustar anchos columnas
    column_widths = [15, 30, 10, 50, 15]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = width

    wb.save("Personajes_Destacados_Con_Imagenes.xlsx")
    print("Archivo Excel creado: Personajes_Destacados_Con_Imagenes.xlsx")
import openpyxl
from openpyxl.styles import Font, Alignment
from openpyxl.drawing.image import Image
import requests
from io import BytesIO

def download_image(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        return BytesIO(response.content)
    except Exception as e:
        print(f"No se pudo descargar la imagen de {url}: {e}")
        return None

def create_excel_with_images():
    data = [
        ('Familia', 'Personaje Destacado', 'Siglo', 'Rol / Contribución', 'Imagen'),
        ('Romero', 'José Romero Silva', 'XVIII', 'Terrateniente, fundador de haciendas', 'https://upload.wikimedia.org/wikipedia/commons/thumb/2/2d/Coat_of_Arms_of_Spain.svg/120px-Coat_of_Arms_of_Spain.svg.png'),
        ('Hernández', 'Manuel Hernández Fuentes', 'XVIII', 'Agricultor destacado en Ñuble', 'https://upload.wikimedia.org/wikipedia/commons/thumb/a/a3/Coat_of_Arms_of_Chile.svg/120px-Coat_of_Arms_of_Chile.svg.png'),
        ('Valdés', 'José Miguel Valdés Hernández', 'XIX', 'Político, empresario agrícola', 'https://upload.wikimedia.org/wikipedia/commons/thumb/a/a3/Coat_of_Arms_of_Chile.svg/120px-Coat_of_Arms_of_Chile.svg.png'),
        ('Fuentes', 'Pedro Fuentes Ríos', 'XVIII', 'Participante en la fundación de ciudades', 'https://upload.wikimedia.org/wikipedia/commons/8/8d/Chilean_Navy_Seal.svg'),
        ('Molina', 'Pedro Molina Fuentes', 'XIX', 'Figura política y médico', 'https://upload.wikimedia.org/wikipedia/commons/7/7e/Seal_of_Chile.svg'),
        ('Del Río / Ríos', 'José Del Río Hernández', 'XIX', 'Empresario y político', 'https://upload.wikimedia.org/wikipedia/commons/6/60/Flag_of_Chile.svg'),
        ('Parra', 'Isabel Parra Molina', 'XX', 'Cantante y figura cultural', 'https://upload.wikimedia.org/wikipedia/commons/6/60/Flag_of_Chile.svg'),
        ('León', 'Miguel León Fuentes', 'XVIII', 'Terrateniente y comerciante', 'https://upload.wikimedia.org/wikipedia/commons/thumb/2/2d/Coat_of_Arms_of_Spain.svg/120px-Coat_of_Arms_of_Spain.svg.png'),
        ('Luco', 'Luis Luco Fuentes', 'XIX', 'Político y empresario agrícola', 'https://upload.wikimedia.org/wikipedia/commons/thumb/a/a3/Coat_of_Arms_of_Chile.svg/120px-Coat_of_Arms_of_Chile.svg.png'),
        ('Aguirre', 'Francisco Aguirre Coloma', 'XVIII', 'Militar y político', 'https://upload.wikimedia.org/wikipedia/commons/8/8d/Chilean_Navy_Seal.svg'),
        ('Coloma', 'Manuel Coloma Hernández', 'XVIII', 'Religioso y educador', 'https://upload.wikimedia.org/wikipedia/commons/7/7e/Seal_of_Chile.svg'),
    ]

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Personajes Destacados"

    for row_index, row_data in enumerate(data, 1):
        for col_index, value in enumerate(row_data, 1):
            if row_index == 1 or col_index != 5:
                # Encabezados o datos normales
                cell = ws.cell(row=row_index, column=col_index, value=value)
                if row_index == 1:
                    cell.font = Font(bold=True)
                    cell.alignment = Alignment(horizontal='center')
                else:
                    cell.alignment = Alignment(wrap_text=True)
            else:
                # Insertar imagen en columna 5 (índice 4)
                img_data = download_image(value)
                if img_data:
                    img = Image(img_data)
                    img.width = 50
                    img.height = 50
                    cell_location = f"E{row_index}"
                    ws.add_image(img, cell_location)
                else:
                    ws.cell(row=row_index, column=col_index, value="Imagen no disponible")

    # Ajustar anchos columnas
    column_widths = [15, 30, 10, 50, 15]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = width

    wb.save("Personajes_Destacados_Con_Imagenes.xlsx")
    print("Archivo Excel creado: Personajes_Destacados_Con_Imagenes.xlsx")
from docx import Document
from docx.shared import Inches
import requests
from io import BytesIO

def download_image(url):
    response = requests.get(url)
    return BytesIO(response.content)

def crear_word_con_imagenes():
    doc = Document()
    doc.add_heading('Personajes Destacados y sus Familias', 0)

    personajes = [
        ('Romero', 'José Romero Silva', 'Terrateniente, fundador de haciendas', 'https://upload.wikimedia.org/wikipedia/commons/thumb/2/2d/Coat_of_Arms_of_Spain.svg/120px-Coat_of_Arms_of_Spain.svg.png'),
        ('Hernández', 'Manuel Hernández Fuentes', 'Agricultor destacado en Ñuble', 'https://upload.wikimedia.org/wikipedia/commons/thumb/a/a3/Coat_of_Arms_of_Chile.svg/120px-Coat_of_Arms_of_Chile.svg.png'),
        # ... añade más personajes aquí
    ]

    for familia, nombre, rol, img_url in personajes:
        doc.add_heading(f'{nombre} ({familia})', level=2)
        doc.add_paragraph(rol)
        try:
            img_stream = download_image(img_url)
            doc.add_picture(img_stream, width=Inches(1.0))
        except Exception as e:
            doc.add_paragraph("Imagen no disponible.")
    
    doc.save('Personajes_Con_Imagenes.docx')
    print("Archivo Word creado: Personajes_Con_Imagenes.docx")

if __name__ == "__main__":
    crear_word_con_imagenes()

if __name__ == "__main__":
    create_excel_with_images()
    import openpyxl
from openpyxl.styles import Font, Alignment
from openpyxl.drawing.image import Image
import requests
from io import BytesIO
from docx import Document
from docx.shared import Inches

# Datos base
personajes = [
    ('Romero', 'José Romero Silva', 'XVIII', 'Terrateniente, fundador de haciendas', 'https://upload.wikimedia.org/wikipedia/commons/thumb/2/2d/Coat_of_Arms_of_Spain.svg/120px-Coat_of_Arms_of_Spain.svg.png'),
    ('Hernández', 'Manuel Hernández Fuentes', 'XVIII', 'Agricultor destacado en Ñuble', 'https://upload.wikimedia.org/wikipedia/commons/thumb/a/a3/Coat_of_Arms_of_Chile.svg/120px-Coat_of_Arms_of_Chile.svg.png'),
    ('Valdés', 'José Miguel Valdés Hernández', 'XIX', 'Político, empresario agrícola', 'https://upload.wikimedia.org/wikipedia/commons/thumb/a/a3/Coat_of_Arms_of_Chile.svg/120px-Coat_of_Arms_of_Chile.svg.png'),
    ('Fuentes', 'Pedro Fuentes Ríos', 'XVIII', 'Participante en la fundación de ciudades', 'https://upload.wikimedia.org/wikipedia/commons/8/8d/Chilean_Navy_Seal.svg'),
    ('Molina', 'Pedro Molina Fuentes', 'XIX', 'Figura política y médico', 'https://upload.wikimedia.org/wikipedia/commons/7/7e/Seal_of_Chile.svg'),
    ('Del Río / Ríos', 'José Del Río Hernández', 'XIX', 'Empresario y político', 'https://upload.wikimedia.org/wikipedia/commons/6/60/Flag_of_Chile.svg'),
    ('Parra', 'Isabel Parra Molina', 'XX', 'Cantante y figura cultural', 'https://upload.wikimedia.org/wikipedia/commons/6/60/Flag_of_Chile.svg'),
    ('León', 'Miguel León Fuentes', 'XVIII', 'Terrateniente y comerciante', 'https://upload.wikimedia.org/wikipedia/commons/thumb/2/2d/Coat_of_Arms_of_Spain.svg/120px-Coat_of_Arms_of_Spain.svg.png'),
    ('Luco', 'Luis Luco Fuentes', 'XIX', 'Político y empresario agrícola', 'https://upload.wikimedia.org/wikipedia/commons/thumb/a/a3/Coat_of_Arms_of_Chile.svg/120px-Coat_of_Arms_of_Chile.svg.png'),
    ('Aguirre', 'Francisco Aguirre Coloma', 'XVIII', 'Militar y político', 'https://upload.wikimedia.org/wikipedia/commons/8/8d/Chilean_Navy_Seal.svg'),
    ('Coloma', 'Manuel Coloma Hernández', 'XVIII', 'Religioso y educador', 'https://upload.wikimedia.org/wikipedia/commons/7/7e/Seal_of_Chile.svg'),
]

def download_image(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        return BytesIO(response.content)
    except Exception as e:
        print(f"Error descargando imagen de {url}: {e}")
        return None

def crear_excel(personajes):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Personajes Destacados"

    headers = ['Familia', 'Personaje', 'Siglo', 'Rol / Contribución', 'Imagen']
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')

    for fila, (familia, nombre, siglo, rol, url_img) in enumerate(personajes, 2):
        ws.cell(row=fila, column=1, value=familia)
        ws.cell(row=fila, column=2, value=nombre)
        ws.cell(row=fila, column=3, value=siglo)
        ws.cell(row=fila, column=4, value=rol)

        img_data = download_image(url_img)
        if img_data:
            img = Image(img_data)
            img.width = 40
            img.height = 40
            celda = f"E{fila}"
            ws.add_image(img, celda)
        else:
            ws.cell(row=fila, column=5, value="Imagen no disponible")

    col_widths = [15, 30, 10, 50, 15]
    for i, width in enumerate(col_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = width

    wb.save("Personajes_Destacados.xlsx")
    print("Excel creado: Personajes_Destacados.xlsx")

def crear_word(personajes):
    doc = Document()
    doc.add_heading('Personajes Destacados y sus Familias', 0)

    for familia, nombre, siglo, rol, img_url in personajes:
        doc.add_heading(f'{nombre} ({familia})', level=2)
        doc.add_paragraph(f'Siglo: {siglo}')
        doc.add_paragraph(rol)
        img_stream = download_image(img_url)
        if img_stream:
            doc.add_picture(img_stream, width=Inches(1))
        else:
            doc.add_paragraph("Imagen no disponible.")
        doc.add_paragraph('')

    doc.save('Personajes_Con_Imagenes.docx')
    print("Word creado: Personajes_Con_Imagenes.docx")

def crear_sql(personajes):
    with open('personajes.sql', 'w', encoding='utf-8') as f:
        f.write("-- Creación tabla personajes\n")
        f.write("""
CREATE TABLE personajes (
    id SERIAL PRIMARY KEY,
    familia VARCHAR(50),
    nombre VARCHAR(100),
    siglo VARCHAR(10),
    rol VARCHAR(255),
    imagen_url VARCHAR(255)
);\n\n""")
        for familia, nombre, siglo, rol, img_url in personajes:
            insert = f"INSERT INTO personajes (familia, nombre, siglo, rol, imagen_url) VALUES ('{familia}', '{nombre}', '{siglo}', '{rol}', '{img_url}');\n"
            f.write(insert)
    print("Archivo SQL creado: personajes.sql")

import psycopg2

def obtener_personajes():
    try:
        conn = psycopg2.connect(
            host="localhost",
            database="familias_chile",
            user="tu_usuario",
            password="tu_contraseña"
        )
        cursor = conn.cursor()

        cursor.execute("SELECT familia, nombre, siglo, rol, imagen_url FROM personajes ORDER BY siglo")
        personajes = cursor.fetchall()

        for fam, nom, siglo, rol, url in personajes:
            print(f"{nom} ({fam}) - Siglo: {siglo} - Rol: {rol}\nImagen: {url}\n")

        cursor.close()
        conn.close()

    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    obtener_personajes()

if __name__ == "__main__":
    crear_excel(personajes)
    crear_word(personajes)
    crear_sql(personajes)

