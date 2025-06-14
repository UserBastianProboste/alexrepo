
import openpyxl
from openpyxl.styles import Font, Alignment

def create_excel():
    data = [
        ('Familia', 'Personaje Destacado', 'Siglo', 'Rol / Contribución'),
        ('Romero', 'José Romero Silva', 'XVIII', 'Terrateniente, fundador de haciendas'),
        ('Romero', 'Pedro Romero Valdés', 'XIX', 'Diputado, empresario agrícola'),
        ('Hernández', 'Manuel Hernández Fuentes', 'XVIII', 'Agricultor destacado en Ñuble'),
        ('Hernández', 'José Hernández Parra', 'XIX', 'Político local, participante independencia'),
        ('Valdés', 'José Miguel Valdés Hernández', 'XIX', 'Político, empresario agrícola'),
        ('Valdés', 'Francisco Valdés Coloma', 'XX', 'Militar y diplomático'),
        ('Fuentes', 'Pedro Fuentes Ríos', 'XVIII', 'Educador y funcionario público'),
        ('Fuentes', 'Juan Fuentes Aguirre', 'XIX', 'Empresario y hacendado'),
        ('Molina', 'Pedro Molina Fuentes', 'XIX', 'Educador y promotor cultural'),
        ('Molina', 'María Molina Parra', 'XX', 'Activista social'),
        ('Del Río / Ríos', 'José Del Río Hernández', 'XIX', 'Empresario agrícola y político regional'),
        ('Del Río / Ríos', 'Luis Ríos Aguirre', 'XX', 'Militar y político'),
        ('Parra', 'Isabel Parra Molina', 'XX', 'Educadora y promotora cultural'),
        ('Parra', 'Carlos Parra Fuentes', 'XIX', 'Empresario'),
        ('León', 'Miguel León Fuentes', 'XIX', 'Terrateniente y político local'),
        ('Luco', 'Luis Luco Fuentes', 'XIX', 'Diputado y abogado'),
        ('Aguirre', 'Francisco Aguirre Coloma', 'XIX', 'Militar y político'),
        ('Coloma', 'Manuel Coloma Hernández', 'XIX', 'Terrateniente y político'),
    ]

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Personajes Destacados"

    # Agregar datos
    for row_index, row_data in enumerate(data, 1):
        for col_index, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_index, column=col_index, value=value)
            # Encabezado en negrita
            if row_index == 1:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal='center')
            else:
                cell.alignment = Alignment(wrap_text=True)

    # Ajustar ancho columnas
    column_widths = [15, 30, 10, 50]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = width

    wb.save("Personajes_Destacados.xlsx")
    print("Archivo Excel creado: Personajes_Destacados.xlsx")
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_header_footer(doc, author_name):
    # Agregar encabezado
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = (
        "Familias Históricas y Genealogía en Chile\n"
        "De los Siglos XV al XXI\n"
        f"Documento elaborado por {author_name}\n"
        "Fecha: 14 de junio de 2025"
    )
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_para.runs[0].font.size = Pt(10)
    header_para.runs[0].font.name = 'Calibri'

    # Agregar pie de página
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"© 2025 {author_name} – Investigación Genealógica"
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.name = 'Calibri'

    # Numeración de página en pie de página (al centro)
    # Crear un elemento de campo PAGE
    page_num_run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')  # start
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')  # separate
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')  # end
    fldChar3.set(qn('w:fldCharType'), 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)
    page_num_run._r.append(fldChar3)

    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def create_personajes_table(doc):
    # Tabla con personajes destacados
    rows = 20
    cols = 4
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    # Encabezados
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Familia'
    hdr_cells[1].text = 'Personaje Destacado'
    hdr_cells[2].text = 'Siglo'
    hdr_cells[3].text = 'Rol / Contribución'

    # Datos a insertar
    data = [
        ('Romero', 'José Romero Silva', 'XVIII', 'Terrateniente, fundador de haciendas'),
        ('Romero', 'Pedro Romero Valdés', 'XIX', 'Diputado, empresario agrícola'),
        ('Hernández', 'Manuel Hernández Fuentes', 'XVIII', 'Agricultor destacado en Ñuble'),
        ('Hernández', 'José Hernández Parra', 'XIX', 'Político local, participante independencia'),
        ('Valdés', 'José Miguel Valdés Hernández', 'XIX', 'Político, empresario agrícola'),
        ('Valdés', 'Francisco Valdés Coloma', 'XX', 'Militar y diplomático'),
        ('Fuentes', 'Pedro Fuentes Ríos', 'XVIII', 'Educador y funcionario público'),
        ('Fuentes', 'Juan Fuentes Aguirre', 'XIX', 'Empresario y hacendado'),
        ('Molina', 'Pedro Molina Fuentes', 'XIX', 'Educador y promotor cultural'),
        ('Molina', 'María Molina Parra', 'XX', 'Activista social'),
        ('Del Río / Ríos', 'José Del Río Hernández', 'XIX', 'Empresario agrícola y político regional'),
        ('Del Río / Ríos', 'Luis Ríos Aguirre', 'XX', 'Militar y político'),
        ('Parra', 'Isabel Parra Molina', 'XX', 'Educadora y promotora cultural'),
        ('Parra', 'Carlos Parra Fuentes', 'XIX', 'Empresario'),
        ('León', 'Miguel León Fuentes', 'XIX', 'Terrateniente y político local'),
        ('Luco', 'Luis Luco Fuentes', 'XIX', 'Diputado y abogado'),
        ('Aguirre', 'Francisco Aguirre Coloma', 'XIX', 'Militar y político'),
        ('Coloma', 'Manuel Coloma Hernández', 'XIX', 'Terrateniente y político'),
    ]

    # Llenar filas
    for i, row_data in enumerate(data, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row_data):
            cells[j].text = val

    # Ajustar estilo fuente y tamaño
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = 'Normal'
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

def create_intro(doc):
    intro_text = (
        "Este documento contiene un resumen genealógico y de personajes destacados de las "
        "familias históricas con influencia en Chile desde el siglo XV hasta la actualidad. "
        "Incluye información relevante sobre la genealogía, política, militar, social y cultural."
    )
    p = doc.add_paragraph(intro_text)
    p.style = 'Normal'
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

def main():
    author_name = "Alex Romero"  # Cambia esto por tu nombre si quieres
    doc = Document()

    # Crear encabezado y pie de página
    add_header_footer(doc, author_name)

    # Título principal
    titulo = doc.add_heading('Familias Históricas y Genealogía en Chile', level=1)
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # espacio

    # Introducción
    create_intro(doc)

    doc.add_paragraph()  # espacio

    # Tabla personajes
    doc.add_heading('Personajes Destacados por Familia', level=2)
    create_personajes_table(doc)

    # Guardar archivo
    doc.save('Familias_Genealogia_Chile.docx')
    print("Documento creado: Familias_Genealogia_Chile.docx")
    from pymongo import MongoClient
from docx import Document
from docx.shared import Pt

client = MongoClient('mongodb://localhost:27017/')
db = client['familias_chile']
coleccion = db['genealogia']

def print_genealogia_doc(persona, doc, nivel=0):
    indent = "  " * nivel
    p = doc.add_paragraph()
    run = p.add_run(f"{indent}{persona['nombre']} ({persona['familia']}) - Siglo: {persona['siglo']} - Rol: {persona['rol']}")
    run.font.size = Pt(12 - nivel)  # Fuente más chica según nivel
    for hijo in persona.get("descendientes", []):
        print_genealogia_doc(hijo, doc, nivel + 1)

def generar_word():
    doc = Document()
    doc.add_heading('Genealogía Familias Chile', 0)

    for persona in coleccion.find():
        print_genealogia_doc(persona, doc)
        doc.add_paragraph()  # Espacio entre árboles

    nombre_archivo = "Genealogia_Familias_Chile.docx"
    doc.save(nombre_archivo)
    print(f"Documento guardado: {nombre_archivo}")

if __name__ == "__main__":
    generar_word()

if __name__ == "__main__":
    create_excel()