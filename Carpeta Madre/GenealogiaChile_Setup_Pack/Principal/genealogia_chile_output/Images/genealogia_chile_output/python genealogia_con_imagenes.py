import os
import sqlite3
import pandas as pd
from docx import Document
from fpdf import FPDF
import requests
from PIL import Image
from io import BytesIO
from openpyxl import load_workbook
from openpyxl.drawing.image import Image as XLImage

# Carpeta de salida
output_dir = "output"
os.makedirs(output_dir, exist_ok=True)

# Datos base
familias = [
    {"apellido": "Romero", "origen": "España", "contribucion": "Ejército Libertador", "img_url": "https://upload.wikimedia.org/wikipedia/commons/thumb/8/8f/Coat_of_Arms_of_Spain_(1785-1873_and_1875-1931).svg/320px-Coat_of_Arms_of_Spain_(1785-1873_and_1875-1931).svg.png"},
    {"apellido": "Hernández", "origen": "España", "contribucion": "Constitución y colonización", "img_url": "https://upload.wikimedia.org/wikipedia/commons/thumb/f/f3/Flag_of_Spain.svg/2560px-Flag_of_Spain.svg.png"},
    {"apellido": "Valdés", "origen": "España", "contribucion": "Fundación de ciudades", "img_url": "https://upload.wikimedia.org/wikipedia/commons/thumb/f/f5/Flag_of_Chile.svg/2560px-Flag_of_Chile.svg.png"},
    {"apellido": "Fuentes", "origen": "España", "contribucion": "Política del siglo XIX", "img_url": "https://upload.wikimedia.org/wikipedia/commons/6/62/Coat_of_arms_of_Chile.svg"},
]

# ========================
# Crear Documento Word
# ========================
doc = Document()
doc.add_heading("Resumen Genealógico de Familias Históricas", 0)
for fam in familias:
    doc.add_heading(fam["apellido"], level=1)
    doc.add_paragraph(f'Origen: {fam["origen"]}')
    doc.add_paragraph(f'Contribución histórica: {fam["contribucion"]}')
doc.save(os.path.join(output_dir, "resumen_genealogico.docx"))

# ========================
# Crear Documento PDF
# ========================
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)
pdf.cell(200, 10, txt="Resumen Genealógico de Familias Históricas", ln=True, align="C")
pdf.ln(10)
for fam in familias:
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(200, 10, txt=fam["apellido"], ln=True)
    pdf.set_font("Arial", size=11)
    pdf.cell(200, 8, txt=f"Origen: {fam['origen']}", ln=True)
    pdf.cell(200, 8, txt=f"Contribución: {fam['contribucion']}", ln=True)
    pdf.ln(4)
pdf.output(os.path.join(output_dir, "resumen_genealogico.pdf"))

# ========================
# Crear Base de Datos SQLite
# ========================
conn = sqlite3.connect(os.path.join(output_dir, "genealogia.db"))
cursor = conn.cursor()
cursor.execute("DROP TABLE IF EXISTS Familias")
cursor.execute("""
    CREATE TABLE Familias (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        apellido TEXT,
        origen TEXT,
        contribucion TEXT
    )
""")
for fam in familias:
    cursor.execute("INSERT INTO Familias (apellido, origen, contribucion) VALUES (?, ?, ?)",
                   (fam["apellido"], fam["origen"], fam["contribucion"]))
conn.commit()
conn.close()

# ========================
# Crear Excel con imágenes
# ========================
excel_path = os.path.join(output_dir, "familias_genealogicas.xlsx")
df = pd.DataFrame(familias).drop(columns=["img_url"])
df.to_excel(excel_path, index=False)

# Descargar e insertar imágenes
wb = load_workbook(excel_path)
ws = wb.active

for idx, fam in enumerate(familias, start=2):
    img_url = fam["img_url"]
    try:
        response = requests.get(img_url)
        image = Image.open(BytesIO(response.content))
        image_path = os.path.join(output_dir, f"{fam['apellido']}_img.png")
        image.save(image_path)

        img = XLImage(image_path)
        img.width = 60
        img.height = 40
        cell = f"E{idx}"  # Columna E
        ws.add_image(img, cell)

    except Exception as e:
        print(f"⚠️ Error al cargar imagen de {fam['apellido']}: {e}")

# Guardar Excel con imágenes
wb.save(excel_path)

print("✅ ¡Todos los archivos han sido generados correctamente!")
