import sqlite3
import os
from docx import Document
import sys
import os
import sqlite3
create-react-app genealogia-chile
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QTableWidget, QTableWidgetItem, QMessageBox,
    QHeaderView, QAbstractItemView, QScrollArea
)
from PyQt6.QtGui import QPixmap
from PyQt6.QtCore import Qt
from docx import Document
from fpdf import FPDF
import pandas as pd
from openpyxl import load_workbook
import requests
from PIL import Image
from io import BytesIO

# Carpeta output
OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Base de datos
DB_PATH = os.path.join(OUTPUT_DIR, "genealogia.db")

# Datos de ejemplo (mismo que antes)
familias = [
    {"apellido": "Romero", "origen": "España", "contribucion": "Ejército Libertador", "img_url": "https://upload.wikimedia.org/wikipedia/commons/thumb/8/8f/Coat_of_Arms_of_Spain_(1785-1873_and_1875-1931).svg/320px-Coat_of_Arms_of_Spain_(1785-1873_and_1875-1931).svg.png"},
    {"apellido": "Hernández", "origen": "España", "contribucion": "Constitución y colonización", "img_url": "https://upload.wikimedia.org/wikipedia/commons/thumb/f/f3/Flag_of_Spain.svg/2560px-Flag_of_Spain.svg.png"},
    {"apellido": "Valdés", "origen": "España", "contribucion": "Fundación de ciudades", "img_url": "https://upload.wikimedia.org/wikipedia/commons/thumb/f/f5/Flag_of_Chile.svg/2560px-Flag_of_Chile.svg.png"},
    {"apellido": "Fuentes", "origen": "España", "contribucion": "Política del siglo XIX", "img_url": "https://upload.wikimedia.org/wikipedia/commons/6/62/Coat_of_arms_of_Chile.svg"},
]
import sqlite3
import os
from docx import Document

OUTPUT_DIR = "genealogia_chile_output"
os.makedirs(OUTPUT_DIR, exist_ok=True)
DB_PATH = os.path.join(OUTPUT_DIR, "genealogia_familias.db")
WORD_PATH = os.path.join(OUTPUT_DIR, "arbol_genealogico.docx")

def crear_base_datos():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''
    CREATE TABLE IF NOT EXISTS persona (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nombre TEXT NOT NULL,
        familia TEXT NOT NULL,
        siglo TEXT,
        rol TEXT,
        padre_id INTEGER,
        FOREIGN KEY(padre_id) REFERENCES persona(id)
    )
    ''')
    conn.commit()
    conn.close()

def insertar_personas():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    personas = [
        (1, "José Miguel Valdés", "Valdés", "XIX", "Político y empresario", None),
        (2, "Ramón Valdés", "Valdés", "XX", "Político", 1),
        (3, "María Valdés", "Valdés", "XX", "Empresaria", 1),
        (4, "Juan Hernández", "Hernández", "XVIII", "Primer chileno de la familia", None),
        (5, "Pedro Hernández", "Hernández", "XIX", "Militar", 4),
        (6, "Luis Hernández", "Hernández", "XX", "Político", 5),
        (7, "Manuel Molina", "Molina", "XIX", "Empresario agrícola", None),
        (8, "Ana Molina", "Molina", "XX", "Educadora", 7),
        (9, "Carlos Parra", "Parra", "XIX", "Abogado", None),
        (10, "Isabel Parra", "Parra", "XX", "Artista", 9),
        (11, "Francisco Luco", "Luco", "XVIII", "Fundador de hacienda", None),
        (12, "José Fuentes", "Fuentes", "XIX", "Político local", None),
        (13, "María Fuentes", "Fuentes", "XX", "Profesora", 12),
        (14, "Ricardo Aguirre", "Aguirre", "XIX", "Militar", None),
        (15, "Sofía Coloma", "Coloma", "XX", "Religiosa", None),
        # Puedes agregar más aquí
    ]
    c.execute("DELETE FROM persona")
    for p in personas:
        c.execute('''
            INSERT INTO persona (id, nombre, familia, siglo, rol, padre_id) VALUES (?, ?, ?, ?, ?, ?)
        ''', p)
    conn.commit()
    conn.close()

def obtener_personas():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT id, nombre, familia, siglo, rol, padre_id FROM persona")
    personas = c.fetchall()
    conn.close()
    return personas

def construir_arbol(personas):
    personas_dict = {p[0]: {'nombre': p[1], 'familia': p[2], 'siglo': p[3], 'rol': p[4], 'padre_id': p[5], 'hijos': []} for p in personas}
    raiz = []
    for pid, data in personas_dict.items():
        padre_id = data['padre_id']
        if padre_id and padre_id in personas_dict:
            personas_dict[padre_id]['hijos'].append(pid)
        else:
            raiz.append(pid)
    return personas_dict, raiz

def agregar_persona_doc(doc, personas_dict, nodos, nivel=0):
    for nodo in nodos:
        persona = personas_dict[nodo]
        indent = "    " * nivel
        texto = f"{indent}- {persona['nombre']} ({persona['familia']}, {persona['siglo']}) - {persona['rol']}"
        doc.add_paragraph(texto)
        if persona['hijos']:
            agregar_persona_doc(doc, personas_dict, persona['hijos'], nivel+1)

def generar_word():
    personas = obtener_personas()
    personas_dict, raiz = construir_arbol(personas)
    doc = Document()
    doc.add_heading('Árbol Genealógico de Familias Históricas de Chile', 0)
    agregar_persona_doc(doc, personas_dict, raiz)
    doc.save(WORD_PATH)
    print(f"[WORD] Archivo guardado en {WORD_PATH}")

def main():
    crear_base_datos()
    insertar_personas()
    generar_word()

if __name__ == "__main__":
    main()

def init_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("DROP TABLE IF EXISTS Familias")
    cursor.execute("""
        CREATE TABLE Familias (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            apellido TEXT,
            origen TEXT,
            contribucion TEXT,
            img_url TEXT
        )
    """)
    for fam in familias:
        cursor.execute("INSERT INTO Familias (apellido, origen, contribucion, img_url) VALUES (?, ?, ?, ?)",
                       (fam["apellido"], fam["origen"], fam["contribucion"], fam["img_url"]))
    conn.commit()
    conn.close()

def fetch_familias():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT id, apellido, origen, contribucion, img_url FROM Familias")
    rows = cursor.fetchall()
    conn.close()
    return rows

def generate_word():
    doc = Document()
    doc.add_heading("Resumen Genealógico de Familias Históricas", 0)
    for fam in familias:
        doc.add_heading(fam["apellido"], level=1)
        doc.add_paragraph(f'Origen: {fam["origen"]}')
        doc.add_paragraph(f'Contribución histórica: {fam["contribucion"]}')
    path = os.path.join(OUTPUT_DIR, "resumen_genealogico.docx")
    doc.save(path)
    return path

def generate_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Resumen Genealógico de Familias Históricas", ln=True, align="C")
    pdf.ln(10)
    for fam in familias:
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(200, 10, txt=fam["apellido"], ln=True)
        pdf.set_font("Arial", size=11)
        pdf.cell(200, 8, txt=f"Origen: {fam['origen']}", ln=True)
        pdf.cell(200, 8, txt=f"Contribución: {fam['contribucion']}", ln=True)
        pdf.ln(4)
    path = os.path.join(OUTPUT_DIR, "resumen_genealogico.pdf")
    pdf.output(path)
    return path

def generate_excel():
    import pandas as pd
    from openpyxl import load_workbook
    from openpyxl.drawing.image import Image as XLImage

    excel_path = os.path.join(OUTPUT_DIR, "familias_genealogicas.xlsx")
    df = pd.DataFrame(familias).drop(columns=["img_url"])
    df.to_excel(excel_path, index=False)

    wb = load_workbook(excel_path)
    ws = wb.active

    for idx, fam in enumerate(familias, start=2):
        img_url = fam["img_url"]
        try:
            response = requests.get(img_url)
            image = Image.open(BytesIO(response.content))
            image_path = os.path.join(OUTPUT_DIR, f"{fam['apellido']}_img.png")
            image.save(image_path)

            img = XLImage(image_path)
            img.width = 60
            img.height = 40
            cell = f"E{idx}"  # Columna E
            ws.add_image(img, cell)

        except Exception as e:
            print(f"⚠️ Error al cargar imagen de {fam['apellido']}: {e}")

    wb.save(excel_path)
    return excel_path

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Genealogía Familias Históricas - Chile")
        self.resize(900, 600)
        self.initUI()

    def initUI(self):
        widget = QWidget()
        self.setCentralWidget(widget)
        layout = QHBoxLayout()
        widget.setLayout(layout)

        # Tabla
        self.table = QTableWidget()
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(["ID", "Apellido", "Origen", "Contribución"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        self.table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.table.cellClicked.connect(self.show_details)

        layout.addWidget(self.table, 70)

        # Panel derecho (detalles e imagen)
        right_panel = QVBoxLayout()

        self.lbl_title = QLabel("Selecciona una familia")
        self.lbl_title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.lbl_title.setStyleSheet("font-weight: bold; font-size: 18px;")
        right_panel.addWidget(self.lbl_title)

        self.lbl_details = QLabel("")
        self.lbl_details.setWordWrap(True)
        right_panel.addWidget(self.lbl_details)

        self.img_label = QLabel()
        self.img_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.img_label.setFixedSize(250, 150)
        right_panel.addWidget(self.img_label)

        # Botones exportar
        btn_layout = QHBoxLayout()
        self.btn_word = QPushButton("Exportar Word")
        self.btn_pdf = QPushButton("Exportar PDF")
        self.btn_excel = QPushButton("Exportar Excel")

        self.btn_word.clicked.connect(self.export_word)
        self.btn_pdf.clicked.connect(self.export_pdf)
        self.btn_excel.clicked.connect(self.export_excel)

        btn_layout.addWidget(self.btn_word)
        btn_layout.addWidget(self.btn_pdf)
        btn_layout.addWidget(self.btn_excel)

        right_panel.addLayout(btn_layout)

        layout.addLayout(right_panel, 30)

        self.load_data()

    def load_data(self):
        self.table.setRowCount(0)
        rows = fetch_familias()
        for row_idx, (fid, apellido, origen, contribucion, img_url) in enumerate(rows):
            self.table.insertRow(row_idx)
            self.table.setItem(row_idx, 0, QTableWidgetItem(str(fid)))
            self.table.setItem(row_idx, 1, QTableWidgetItem(apellido))
            self.table.setItem(row_idx, 2, QTableWidgetItem(origen))
            self.table.setItem(row_idx, 3, QTableWidgetItem(contribucion))

    def show_details(self, row, column):
        fid = self.table.item(row, 0).text()
        apellido = self.table.item(row, 1).text()
        origen = self.table.item(row, 2).text()
        contribucion = self.table.item(row, 3).text()

        self.lbl_title.setText(familias[int(fid)-1]["apellido"])
        self.lbl_details.setText(f"Origen: {origen}\nContribución histórica:\n{contribucion}")

        # Cargar imagen
        img_url = familias[int(fid)-1]["img_url"]
        try:
            response = requests.get(img_url)
            pixmap = QPixmap()
            pixmap.loadFromData(response.content)
            scaled = pixmap.scaled(self.img_label.size(), Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
            self.img_label.setPixmap(scaled)
        except Exception as e:
            self.img_label.clear()

    def export_word(self):
        path = generate_word()
        QMessageBox.information(self, "Exportar Word", f"Archivo guardado en:\n{path}")

    def export_pdf(self):
        path = generate_pdf()
        QMessageBox.information(self, "Exportar PDF", f"Archivo guardado en:\n{path}")

    def export_excel(self):
        path = generate_excel()
        QMessageBox.information(self, "Exportar Excel", f"Archivo guardado en:\n{path}")

def main():
    init_db()  # Crear DB y datos si no existen
    app = QApplication(sys.argv)
    win = MainWindow()
    win.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()

OUTPUT_DIR = "genealogia_chile_output"
os.makedirs(OUTPUT_DIR, exist_ok=True)
DB_PATH = os.path.join(OUTPUT_DIR, "genealogia_familias.db")
WORD_PATH = os.path.join(OUTPUT_DIR, "arbol_genealogico.docx")

def crear_base_datos():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''
    CREATE TABLE IF NOT EXISTS persona (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nombre TEXT NOT NULL,
        familia TEXT NOT NULL,
        siglo TEXT,
        rol TEXT,
        padre_id INTEGER,
        FOREIGN KEY(padre_id) REFERENCES persona(id)
    )
    ''')
    conn.commit()
    conn.close()

def insertar_personas():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    personas = [
        (1, "José Miguel Valdés", "Valdés", "XIX", "Político y empresario", None),
        (2, "Ramón Valdés", "Valdés", "XX", "Político", 1),
        (3, "María Valdés", "Valdés", "XX", "Empresaria", 1),
        (4, "Juan Hernández", "Hernández", "XVIII", "Primer chileno de la familia", None),
        (5, "Pedro Hernández", "Hernández", "XIX", "Militar", 4),
        (6, "Luis Hernández", "Hernández", "XX", "Político", 5),
        (7, "Manuel Molina", "Molina", "XIX", "Empresario agrícola", None),
        (8, "Ana Molina", "Molina", "XX", "Educadora", 7),
        (9, "Carlos Parra", "Parra", "XIX", "Abogado", None),
        (10, "Isabel Parra", "Parra", "XX", "Artista", 9),
        (11, "Francisco Luco", "Luco", "XVIII", "Fundador de hacienda", None),
        (12, "José Fuentes", "Fuentes", "XIX", "Político local", None),
        (13, "María Fuentes", "Fuentes", "XX", "Profesora", 12),
        (14, "Ricardo Aguirre", "Aguirre", "XIX", "Militar", None),
        (15, "Sofía Coloma", "Coloma", "XX", "Religiosa", None),
        # Puedes agregar más aquí
    ]
    c.execute("DELETE FROM persona")
    for p in personas:
        c.execute('''
            INSERT INTO persona (id, nombre, familia, siglo, rol, padre_id) VALUES (?, ?, ?, ?, ?, ?)
        ''', p)
    conn.commit()
    conn.close()

def obtener_personas():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT id, nombre, familia, siglo, rol, padre_id FROM persona")
    personas = c.fetchall()
    conn.close()
    return personas

def construir_arbol(personas):
    personas_dict = {p[0]: {'nombre': p[1], 'familia': p[2], 'siglo': p[3], 'rol': p[4], 'padre_id': p[5], 'hijos': []} for p in personas}
    raiz = []
    for pid, data in personas_dict.items():
        padre_id = data['padre_id']
        if padre_id and padre_id in personas_dict:
            personas_dict[padre_id]['hijos'].append(pid)
        else:
            raiz.append(pid)
    return personas_dict, raiz

def agregar_persona_doc(doc, personas_dict, nodos, nivel=0):
    for nodo in nodos:
        persona = personas_dict[nodo]
        indent = "    " * nivel
        texto = f"{indent}- {persona['nombre']} ({persona['familia']}, {persona['siglo']}) - {persona['rol']}"
        doc.add_paragraph(texto)
        if persona['hijos']:
            agregar_persona_doc(doc, personas_dict, persona['hijos'], nivel+1)

def generar_word():
    personas = obtener_personas()
    personas_dict, raiz = construir_arbol(personas)
    doc = Document()
    doc.add_heading('Árbol Genealógico de Familias Históricas de Chile', 0)
    agregar_persona_doc(doc, personas_dict, raiz)
    doc.save(WORD_PATH)
    print(f"[WORD] Archivo guardado en {WORD_PATH}")

def main():
    insertar_personas()
    insertar_genealogia()
    mostrar_arbol()
    mostrar_arbol()
    generador_genealogia()
    generar_word()
    crear_db_sqlite()

    exportar_excel()


if __name__ == "__main__":
    main()
