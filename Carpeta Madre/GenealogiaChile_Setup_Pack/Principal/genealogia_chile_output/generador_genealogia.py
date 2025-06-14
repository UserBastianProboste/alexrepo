from docx import Document
from fpdf import FPDF
import sqlite3
import pandas as pd

# Datos resumidos de familias históricas
familias = [
    {apellido Romero, origen España, contribucion Ejército Libertador},
    {apellido Hernández, origen España, contribucion Constitución y colonización},
    {apellido Valdés, origen España, contribucion Fundación de ciudades},
    {apellido Fuentes, origen España, contribucion Política del siglo XIX},
    {apellido Molina, origen España, contribucion Educación e Iglesia},
    {apellido Parra, origen España, contribucion Cultura y arte},
    {apellido Luco, origen España, contribucion Justicia y derecho},
    {apellido Coloma, origen España, contribucion Milicia},
    {apellido Silva, origen España, contribucion Gobierno},
    {apellido Aguirre, origen España, contribucion Reforma Agraria}
]

# ========================
# GENERAR WORD
# ========================
doc = Document()
doc.add_heading(Resumen Genealógico de Familias Históricas, 0)

for fam in familias
    doc.add_heading(fam[apellido], level=1)
    doc.add_paragraph(f'Origen {fam[origen]}')
    doc.add_paragraph(f'Contribución histórica {fam[contribucion]}')
    doc.add_paragraph()

doc.save(resumen_genealogico.docx)
print(✅ Documento Word generado resumen_genealogico.docx)

# ========================
# GENERAR PDF
# ========================
pdf = FPDF()
pdf.add_page()
pdf.set_font(Arial, size=12)
pdf.cell(200, 10, txt=Resumen Genealógico de Familias Históricas, ln=True, align=C)
pdf.ln(10)

for fam in familias
    pdf.set_font(Arial, 'B', 12)
    pdf.cell(200, 10, txt=fam[apellido], ln=True)
    pdf.set_font(Arial, size=11)
    pdf.cell(200, 8, txt=fOrigen {fam['origen']}, ln=True)
    pdf.cell(200, 8, txt=fContribución {fam['contribucion']}, ln=True)
    pdf.ln(4)

pdf.output(resumen_genealogico.pdf)
print(✅ Documento PDF generado resumen_genealogico.pdf)

# ========================
# CREAR BASE DE DATOS SQLite
# ========================
conn = sqlite3.connect(genealogia.db)
cursor = conn.cursor()
cursor.execute(DROP TABLE IF EXISTS Familias)
cursor.execute(
    CREATE TABLE Familias (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        apellido TEXT,
        origen TEXT,
        contribucion TEXT
    )
)
for fam in familias
    cursor.execute(INSERT INTO Familias (apellido, origen, contribucion) VALUES (, , ),
                   (fam[apellido], fam[origen], fam[contribucion]))
conn.commit()
conn.close()
print(✅ Base de datos SQLite generada genealogia.db)

# ========================
# CREAR EXCEL
# ========================
df = pd.DataFrame(familias)
df.to_excel(familias_genealogicas.xlsx, index=False)
print(✅ Archivo Excel generado familias_genealogicas.xlsx)
