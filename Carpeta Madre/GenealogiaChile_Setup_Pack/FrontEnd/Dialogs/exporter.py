from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.drawing.image import Image
import requests
from io import BytesIO
from backend.models import Persona, Familia

class Exporter:
    def __init__(self):
        pass

    def exportar_a_word(self, personas, familias, filename="genealogia.docx"):
        doc = Document()
        doc.add_heading("Genealogía de Familias Históricas", 0)

        doc.add_heading("Familias", level=1)
        for fam in familias:
            doc.add_heading(fam.apellido, level=2)
            if fam.origen:
                doc.add_paragraph(f"Origen: {fam.origen}")
            if fam.historia:
                doc.add_paragraph(f"Historia: {fam.historia}")
            if fam.miembros:
                doc.add_paragraph("Miembros destacados:")
                for mid in fam.miembros:
                    # Buscar nombre persona (se asume ya resuelto fuera)
                    persona = next((p for p in personas if p.id == mid), None)
                    if persona:
                        doc.add_paragraph(f" - {persona.nombre} {persona.apellido_paterno} {persona.apellido_materno}", style='ListBullet')

        doc.add_heading("Personas", level=1)
        for p in personas:
            nombre_completo = f"{p.nombre} {p.apellido_paterno} {p.apellido_materno}"
            doc.add_heading(nombre_completo, level=2)
            if p.fecha_nacimiento:
                doc.add_paragraph(f"Fecha de nacimiento: {p.fecha_nacimiento}")
            if p.lugar_nacimiento:
                doc.add_paragraph(f"Lugar de nacimiento: {p.lugar_nacimiento}")
            if p.notas:
                doc.add_paragraph(f"Notas: {p.notas}")

        doc.save(filename)

    def exportar_a_excel(self, personas, familias, filename="genealogia.xlsx"):
        wb = Workbook()
        ws_pers = wb.active
        ws_pers.title = "Personas"

        # Encabezados para personas
        ws_pers.append(["ID", "Nombre", "Apellido Paterno", "Apellido Materno", "Fecha Nac.", "Lugar Nac.", "Notas", "Imagen"])

        for p in personas:
            # Buscar imagen en internet (ejemplo sencillo: foto histórica o escudo, aquí con placeholder)
            url_imagen = f"https://via.placeholder.com/50.png?text={p.apellido_paterno[0].upper()}" if p.apellido_paterno else None

            if url_imagen:
                try:
                    response = requests.get(url_imagen)
                    img_file = BytesIO(response.content)
                    img = Image(img_file)
                    # Insertar imagen en celda será manual al guardar, openpyxl solo permite colocarlo encima
                    # Se ubica en columna H, fila actual +1 (porque header fila 1)
                    row_idx = ws_pers.max_row + 1
                    ws_pers.append([p.id, p.nombre, p.apellido_paterno, p.apellido_materno, p.fecha_nacimiento, p.lugar_nacimiento, p.notas])
                    ws_pers.add_image(img, f"H{row_idx}")
                except Exception:
                    ws_pers.append([p.id, p.nombre, p.apellido_paterno, p.apellido_materno, p.fecha_nacimiento, p.lugar_nacimiento, p.notas])
            else:
                ws_pers.append([p.id, p.nombre, p.apellido_paterno, p.apellido_materno, p.fecha_nacimiento, p.lugar_nacimiento, p.notas])

        # Familias
        ws_fam = wb.create_sheet(title="Familias")
        ws_fam.append(["ID", "Apellido", "Origen", "Historia", "Miembros (IDs)"])
        for fam in familias:
            miembros_str = ", ".join(fam.miembros) if fam.miembros else ""
            ws_fam.append([fam.id, fam.apellido, fam.origen, fam.historia, miembros_str])

        wb.save(filename)
